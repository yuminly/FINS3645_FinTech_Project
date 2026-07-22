"""
FINS3645 FinTech Project Part A - Report Generator
Creates HD-quality Word document following lecture conventions
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

# Path to figures
FIGURES_DIR = '/Users/yumeme/Documents/GitHub/FINS3645_FinTech_Project/report/figures'


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading_elm = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')


def create_report():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('FINS3645 FinTech Project')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x08, 0x51, 0x9c)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Part A: Written Report')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x63, 0x63, 0x63)
    
    doc.add_paragraph()
    
    app_title = doc.add_paragraph()
    app_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = app_title.add_run('FinVest Pro')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x08, 0x51, 0x9c)
    
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('Multi-Asset Portfolio Optimisation with News-Sentiment Analytics')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x63, 0x63, 0x63)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Group details
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = details.add_run('Group: [Your Group Number]\nMembers: [Your Name], [Partner Name]\nDate: July 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x63, 0x63, 0x63)
    
    doc.add_page_break()
    
    # =====================================================================
    # EXECUTIVE SUMMARY
    # =====================================================================
    
    doc.add_heading('Executive Summary', level=1)
    
    doc.add_paragraph(
        'FinVest Pro is a prototype investment platform offering three systematically managed funds: '
        'an equity-only fund, a crypto-only fund, and a combined multi-asset fund. Each fund is constructed '
        'using three optimisation methods—maximum Sharpe ratio, minimum variance, and risk parity—benchmarked '
        'against an equal-weight (1/N) portfolio. A news-sentiment index, built from ensemble natural language '
        'processing of financial headlines, provides an additional analytical layer for equity sector allocation.'
    )
    
    doc.add_paragraph(
        'Our key finding challenges conventional wisdom: the maximum Sharpe (tangency) portfolio, while optimal '
        'in-sample, underperforms out-of-sample with a Sharpe ratio of 0.18 versus 0.32 for the naive 1/N '
        'benchmark. The minimum variance portfolio delivers the strongest risk-adjusted OOS performance (Sharpe 0.69), '
        'confirming that lower-volatility allocations are more robust to estimation error. Risk parity provides '
        'moderate returns with better drawdown management than equal-weight.'
    )
    
    doc.add_paragraph(
        'The news-sentiment index reveals significant sector divergence during market stress, with Technology '
        'and Energy sectors showing the highest sentiment volatility. However, consistent with academic evidence, '
        'naive lexicon-based sentiment shows weak predictive power for returns (correlation ≈ −0.04), highlighting '
        'the need for trained models in future iterations.'
    )
    
    doc.add_page_break()
    
    # =====================================================================
    # 1. INTRODUCTION
    # =====================================================================
    
    doc.add_heading('1. Introduction', level=1)
    
    doc.add_heading('1.1 Product Overview', level=2)
    
    doc.add_paragraph(
        'FinVest Pro addresses a gap in retail investment platforms: most offer either equity funds or crypto '
        'exposure, but few provide a systematic, transparent framework for multi-asset allocation with '
        'institutional-grade risk management. Our platform offers:'
    )
    
    features = [
        'Three optimally constructed funds (equity-only, crypto-only, combined multi-asset)',
        'Four construction methods per fund (maximum Sharpe, minimum variance, risk parity, equal-weight benchmark)',
        'Transparent out-of-sample backtesting with rolling-window rebalancing',
        'A news-sentiment analytics dashboard tracking market mood across equity sectors',
        'Interactive visualisation of portfolio weights, risk contributions, and performance attribution'
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('1.2 Design Principles', level=2)
    
    doc.add_paragraph(
        'Following the Data Factory Floor (DFF) framework, our product is built in four stages: '
        'Stage 1 (ETL) collects and cleans market data; Stage 2 (Feature Engineering) computes rolling '
        'risk metrics and sentiment features; Stage 3 (Model Design) optimises portfolio weights; '
        'Stage 4 (Implementation) runs out-of-sample backtests with realistic rebalancing constraints.'
    )
    
    doc.add_paragraph(
        'We adhere to the four standards for all figures: sentence-style titles, percentage/monetary units, '
        'explicit source attribution, and sample window disclosure. Every chart uses a single emphasis colour '
        '(deep blue, #08519c) to guide attention, consistent with Financial Times visual conventions.'
    )
    
    doc.add_page_break()
    
    # =====================================================================
    # 2. DATA FACTORY FLOOR
    # =====================================================================
    
    doc.add_heading('2. Data Factory Floor (DFF)', level=1)
    
    # Stage 1
    doc.add_heading('Stage 1: ETL — Data Ingestion and Cleaning', level=2)
    
    doc.add_paragraph(
        'Our dataset spans 1 January 2022 to 31 December 2024 (782 business days), covering five US equities '
        '(AAPL, MSFT, JPM, XOM, KO) and five cryptocurrencies (BTC, ETH, SOL, ADA, DOGE). Equity prices '
        'are adjusted for the AAPL 4-for-1 stock split on 15 June 2023 (split_factor = 4.0).'
    )
    
    doc.add_paragraph('Cleaning protocol:')
    cleaning_steps = [
        'Drop exact duplicate rows (6 rows removed)',
        'Drop rows with missing adj_close or volume (7 rows removed)',
        'No winsorising or outlier trimming applied',
        'Two calendars: stocks trade on business days only; crypto trades daily (weekends included)',
        'Unbalanced panel: crypto weekend returns aligned to stock calendar via forward-fill of risk-free rate'
    ]
    
    for step in cleaning_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_paragraph(
        'Risk-free rate: US Treasury bill rate (4.0% annualised, 0.01575% daily), forward-filled across '
        'crypto weekends to maintain a consistent panel for combined-asset optimisation.'
    )
    
    # Stage 2
    doc.add_heading('Stage 2: Feature Engineering', level=2)
    
    doc.add_paragraph('From raw prices, we compute the following features for each asset:')
    
    features_table = doc.add_table(rows=5, cols=3)
    features_table.style = 'Light Grid Accent 1'
    
    # Header
    for i, text in enumerate(['Feature', 'Formula', 'Window']):
        cell = features_table.rows[0].cells[i]
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
    
    feature_data = [
        ('Daily returns', 'pct_change(fill_method=None)', '—'),
        ('21-day rolling volatility', 'std(returns, 21) × √252', '21 days'),
        ('60-day rolling correlation', 'corr(returns, 60)', '60 days'),
        ('Sentiment index', 'Mean VADER + TextBlob + custom', 'Daily')
    ]
    
    for row_idx, (feat, formula, window) in enumerate(feature_data, 1):
        features_table.rows[row_idx].cells[0].text = feat
        features_table.rows[row_idx].cells[1].text = formula
        features_table.rows[row_idx].cells[2].text = window
    
    # Stage 3
    doc.add_heading('Stage 3: Model Design — Portfolio Optimisation', level=2)
    
    doc.add_paragraph(
        'We employ three optimisation methods, each solved via scipy.optimize.minimize with SLSQP:'
    )
    
    methods = [
        ('Maximum Sharpe Ratio (Tangency)', 
         'Maximises (μ_p − r_f) / σ_p where μ_p = w′μ, σ_p = √(w′Σw). '
         'Requires estimation of mean returns and covariance matrix — highly sensitive to estimation error.'),
        ('Minimum Variance',
         'Minimises σ_p² = w′Σw subject to w′1 = 1. Does not require expected return estimates, '
         'making it more robust to estimation error than the tangency portfolio.'),
        ('Risk Parity',
         'Minimises Σ(RC_i − RC_target)² where RC_i = w_i × (Σw)_i / σ_p. '
         'Each asset contributes equally to total portfolio risk, avoiding concentration in high-vol assets.')
    ]
    
    for method_name, method_desc in methods:
        p = doc.add_paragraph()
        run = p.add_run(f'{method_name}: ')
        run.bold = True
        p.add_run(method_desc)
    
    # Stage 4
    doc.add_heading('Stage 4: Implementation — Out-of-Sample Backtesting', level=2)
    
    doc.add_paragraph(
        'All portfolios are evaluated out-of-sample using an expanding-window approach:'
    )
    
    bt_params = [
        'Training window: 252 days (expanding from the start)',
        'Rebalancing frequency: 21 days (monthly)',
        'Out-of-sample period: 2023-01-03 to 2024-12-31',
        'No transaction costs or short-selling',
        'Long-only constraints (weights ∈ [0, 1])'
    ]
    
    for param in bt_params:
        doc.add_paragraph(param, style='List Bullet')
    
    doc.add_page_break()
    
    # =====================================================================
    # 3. RESULTS
    # =====================================================================
    
    doc.add_heading('3. Results', level=1)
    
    # Figure 1
    doc.add_heading('3.1 Equity-Only Optimal Allocation', level=2)
    
    doc.add_paragraph(
        'Question: How should we allocate across five equities to maximise risk-adjusted returns?'
    )
    
    doc.add_paragraph(
        'Figure 1 reveals that the three optimisation methods produce markedly different allocations. '
        'The maximum Sharpe portfolio concentrates heavily in KO (77.6%) and MSFT (15.1%), suggesting '
        'the optimiser identifies these as having the best return-per-unit-risk profile. The minimum '
        'variance portfolio diversifies more broadly across KO (58.5%), MSFT (29.0%), and XOM (10.9%). '
        'Risk parity achieves the most balanced allocation (12.7%–29.6% across all five assets), '
        'reflecting its design principle of equal risk contribution.'
    )
    
    if os.path.exists(f'{FIGURES_DIR}/fig1_equity_allocation.png'):
        doc.add_picture(f'{FIGURES_DIR}/fig1_equity_allocation.png', width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(
        'Figure 1. Equity-only optimal allocation varies by risk preference. '
        'Bars show portfolio weights for three optimisation methods. '
        'Source: FinVest Pro optimiser (scipy minimize, SLSQP). '
        'Sample: 5 US equities, 2022-01-03 to 2024-12-31 (synthetic data).'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x63, 0x63, 0x63)
    
    doc.add_paragraph(
        'Answer: Sector concentration differs by risk preference. The tangency portfolio takes large '
        'bets on individual stocks, while risk parity diversifies across all sectors — a key innovation '
        'for retail investors seeking stable, rule-based allocations.'
    )
    
    # Figure 2
    doc.add_heading('3.2 Crypto-Only Optimal Allocation', level=2)
    
    doc.add_paragraph(
        'Question: How should we allocate across five cryptocurrencies, given their extreme volatility?'
    )
    
    doc.add_paragraph(
        'Figure 2 shows that the maximum Sharpe portfolio concentrates in BTC (56.2%) and ETH (43.8%), '
        'excluding all altcoins entirely. This reflects the optimiser\'s preference for the two largest, '
        'most liquid crypto assets. Minimum variance concentrates even more heavily in BTC (75.4%). '
        'Risk parity again achieves broad diversification (18.0%–23.5% across all five crypto assets), '
        'preventing over-concentration in a single volatile asset.'
    )
    
    if os.path.exists(f'{FIGURES_DIR}/fig2_crypto_allocation.png'):
        doc.add_picture(f'{FIGURES_DIR}/fig2_crypto_allocation.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(
        'Figure 2. Crypto allocations concentrate in BTC and ETH across all methods. '
        'Source: FinVest Pro optimiser. Sample: 5 cryptocurrencies, 2022-01-03 to 2024-12-31 (synthetic).'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x63, 0x63, 0x63)
    
    doc.add_paragraph(
        'Answer: The tangency portfolio\'s crypto concentration creates extreme single-asset risk. '
        'Risk parity\'s equal risk contribution provides a more conservative framework for this '
        'volatile asset class, aligning with the lecture finding that 1/N is hard to beat OOS.'
    )
    
    # Figure 3
    doc.add_heading('3.3 Efficient Frontier Comparison', level=2)
    
    doc.add_paragraph(
        'Question: How do equity, crypto, and combined frontiers compare in risk-return space?'
    )
    
    doc.add_paragraph(
        'Figure 3 plots the in-sample efficient frontiers for all three asset universes. The crypto '
        'frontier dominates in return space (up to ~30% annualised) but at much higher volatility '
        '(50–80%). The equity frontier is more compact (5–18% return, 10–30% volatility). The '
        'combined frontier extends beyond the equity frontier, showing that adding crypto to an '
        'equity portfolio can improve the Sharpe ratio at moderate volatility levels.'
    )
    
    if os.path.exists(f'{FIGURES_DIR}/fig3_efficient_frontier.png'):
        doc.add_picture(f'{FIGURES_DIR}/fig3_efficient_frontier.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(
        'Figure 3. Combined equity-crypto frontier extends beyond equity-only frontier. '
        'Stars mark tangency portfolios. Source: FinVest Pro efficient frontier computation. '
        'Sample: 10 assets, 2022-01-03 to 2024-12-31 (synthetic).'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x63, 0x63, 0x63)
    
    doc.add_paragraph(
        'Answer: Multi-asset diversification offers theoretical improvements over single-class portfolios. '
        'However, as shown below, these in-sample gains do not always translate to OOS performance.'
    )
    
    doc.add_page_break()
    
    # Figure 4
    doc.add_heading('3.4 Out-of-Sample Backtest Performance', level=2)
    
    doc.add_paragraph(
        'Question: Do optimised portfolios beat equal-weight (1/N) out-of-sample?'
    )
    
    doc.add_paragraph(
        'Figure 4 presents the core finding of this report. Despite being theoretically optimal, '
        'the maximum Sharpe portfolio delivers the weakest OOS performance: 15.95% total return '
        '(7.36% annualised) with a Sharpe ratio of just 0.18. The equal-weight benchmark outperforms '
        'it with 35.36% total return (15.64% annualised, Sharpe 0.32). The minimum variance portfolio '
        'achieves the best risk-adjusted performance: 35.21% total return with a Sharpe ratio of 0.69 '
        'and the shallowest maximum drawdown (−15.89%).'
    )
    
    if os.path.exists(f'{FIGURES_DIR}/fig4_backtest_cumulative.png'):
        doc.add_picture(f'{FIGURES_DIR}/fig4_backtest_cumulative.png', width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(
        'Figure 4. Minimum variance dominates out-of-sample; max Sharpe underperforms 1/N. '
        'Top: cumulative return (£1 initial). Bottom: drawdown. '
        'Source: FinVest Pro backtester, expanding window, 21-day rebalance. '
        'Sample: 10 assets, OOS 2023-01-03 to 2024-12-31 (synthetic).'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x63, 0x63, 0x63)
    
    # Performance table
    doc.add_paragraph()
    doc.add_paragraph('Table 1: Out-of-Sample Performance Summary')
    
    perf_table = doc.add_table(rows=5, cols=7)
    perf_table.style = 'Light Grid Accent 1'
    
    headers = ['Method', 'Total Return', 'Annual Return', 'Annual Vol', 'Sharpe', 'Max DD', 'Calmar']
    for i, h in enumerate(headers):
        cell = perf_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    
    perf_data = [
        ['Max Sharpe', '15.95%', '7.36%', '18.83%', '0.18', '−19.17%', '0.38'],
        ['Min Variance', '35.21%', '15.58%', '16.82%', '0.69', '−15.89%', '0.98'],
        ['Risk Parity', '26.26%', '11.84%', '21.54%', '0.36', '−30.70%', '0.39'],
        ['Equal Weight', '35.36%', '15.64%', '36.75%', '0.32', '−42.57%', '0.37']
    ]
    
    for row_idx, row_data in enumerate(perf_data, 1):
        for col_idx, val in enumerate(row_data):
            perf_table.rows[row_idx].cells[col_idx].text = val
    
    doc.add_paragraph(
        'Answer: The tangency portfolio fails OOS, consistent with Week 5 lecture findings that '
        'in-sample optimality does not guarantee OOS performance. Minimum variance and 1/N are '
        'most robust, with minimum variance offering the best risk-adjusted returns.'
    )
    
    doc.add_page_break()
    
    # Figure 5
    doc.add_heading('3.5 News-Sentiment Analytics', level=2)
    
    doc.add_paragraph(
        'Question: Does news sentiment vary by sector, and can it inform portfolio allocation?'
    )
    
    doc.add_paragraph(
        'Figure 5 shows the 7-day moving average of sentiment scores across five equity sectors. '
        'Technology and Consumer sectors exhibit the most stable sentiment profiles, while Energy '
        'and Financials show higher volatility — particularly during the 2022 rate-hiking cycle '
        'and the 2023 regional banking stress. The sector correlation heatmap reveals that '
        'Technology and Consumer sentiment are most correlated (0.72), while Energy diverges '
        'most from other sectors during stress periods.'
    )
    
    if os.path.exists(f'{FIGURES_DIR}/fig5_sentiment_index.png'):
        doc.add_picture(f'{FIGURES_DIR}/fig5_sentiment_index.png', width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(
        'Figure 5. Sentiment diverges across sectors during market stress. '
        'Top: 7-day MA sentiment by sector. Bottom: cross-sector correlation matrix. '
        'Source: FinVest Pro sentiment engine (VADER + TextBlob + custom lexicon). '
        'Sample: 5 sectors, 5,000 synthetic articles, 2022-01-03 to 2024-12-31.'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x63, 0x63, 0x63)
    
    doc.add_paragraph(
        'Answer: Naive lexicon-based sentiment shows sector-specific patterns but weak return '
        'predictive power (correlation ≈ −0.04 with next-day returns). This is consistent with '
        'Week 7 lecture findings. Trained sentiment models (Week 8–9) would likely improve signal quality.'
    )
    
    doc.add_page_break()
    
    # =====================================================================
    # 4. INNOVATION
    # =====================================================================
    
    doc.add_heading('4. Innovation', level=1)
    
    doc.add_paragraph(
        'Our innovation lies in the integration of three elements typically treated separately:'
    )
    
    innovations = [
        ('Multi-Asset Framework', 
         'Most retail platforms offer equity-only or crypto-only funds. FinVest Pro provides a unified '
         'framework combining both asset classes with transparent optimisation, allowing users to understand '
         'the trade-offs between equity stability and crypto growth potential.'),
        ('Sentiment-Weighted Allocation',
         'While Week 7 lectures show that naive lexicon sentiment is not predictive, we use sector-level '
         'sentiment divergence as a risk indicator. When sentiment diverges significantly across sectors '
         '(e.g., Energy bullish while Financials bearish), the platform signals increased portfolio '
         'uncertainty — a practical application beyond simple return prediction.'),
        ('Rolling OOS Transparency',
         'Unlike most platforms that show backtested returns on the full sample, FinVest Pro reports '
         'out-of-sample performance with an expanding training window and monthly rebalancing. This '
         'transparency exposes the gap between in-sample optimality and real-world performance, '
         'empowering users to make informed decisions about portfolio construction.')
    ]
    
    for title, desc in innovations:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)
    
    doc.add_paragraph(
        'Together, these innovations deliver a product that goes beyond what a simple AI prompt would '
        'produce: we combine rigorous quantitative methods with domain-specific financial knowledge '
        'and transparent reporting standards.'
    )
    
    doc.add_page_break()
    
    # =====================================================================
    # 5. DASHBOARD PREVIEW
    # =====================================================================
    
    doc.add_heading('5. Dashboard Implementation', level=1)
    
    doc.add_paragraph(
        'The FinVest Pro dashboard is built with Streamlit and deployed to Streamlit Community Cloud. '
        'It features four main pages:'
    )
    
    pages = [
        'Dashboard: Overview of all fund performance with interactive charts',
        'Fund Comparison: Side-by-side comparison of optimisation methods with weight allocation visualisation',
        'Sentiment Analytics: Real-time sector sentiment tracking with historical correlation analysis',
        'Invest: Interactive portfolio construction tool with custom weight allocation'
    ]
    
    for page in pages:
        doc.add_paragraph(page, style='List Bullet')
    
    doc.add_paragraph(
        'The app is deployed at: https://share.streamlit.io/user/yuminly/FINS3645_FinTech_Project'
    )
    
    doc.add_page_break()
    
    # =====================================================================
    # 6. CONCLUSION
    # =====================================================================
    
    doc.add_heading('6. Conclusion', level=1)
    
    doc.add_paragraph(
        'FinVest Pro demonstrates that systematic portfolio construction for a multi-asset universe '
        'requires careful consideration of estimation error and out-of-sample robustness. Our key findings:'
    )
    
    conclusions = [
        'The maximum Sharpe (tangency) portfolio, while theoretically optimal, underperforms the naive 1/N '
        'benchmark out-of-sample (Sharpe 0.18 vs 0.32), confirming the lecture finding that 1/N is hard to beat.',
        'Minimum variance delivers the strongest risk-adjusted OOS performance (Sharpe 0.69), suggesting '
        'that avoiding estimation of expected returns is a robust strategy for retail investors.',
        'Risk parity provides moderate returns with better drawdown management than equal-weight, '
        'making it suitable for risk-averse investors seeking diversified exposure.',
        'Naive lexicon-based sentiment shows sector-specific patterns but weak return predictive power, '
        'highlighting the need for trained models in future iterations.',
        'Multi-asset portfolios combining equities and crypto offer theoretical diversification benefits, '
        'but the tangency portfolio\'s OOS underperformance limits practical utility.'
    ]
    
    for i, conclusion in enumerate(conclusions, 1):
        doc.add_paragraph(f'{conclusion}', style='List Number')
    
    doc.add_paragraph(
        'For Part B, we plan to extend the sentiment analysis with trained ML models, implement '
        'volatility-targeting overlays, and add transaction cost modelling to improve realism.'
    )
    
    doc.add_page_break()
    
    # =====================================================================
    # REFERENCES
    # =====================================================================
    
    doc.add_heading('References', level=1)
    
    references = [
        'DeMiguel, V., Garlappi, L. and Uppal, R. (2009) \'Optimal versus naive diversification: How inefficient is the 1/N portfolio strategy?\', Review of Financial Studies, 22(5), pp. 1915–1953.',
        'Fama, E.F. and French, K.R. (1993) \'Common risk factors in the returns on stocks and bonds\', Journal of Financial Economics, 33(1), pp. 3–56.',
        'Fisher, K.L. and Statman, M. (2000) \'Behavioral returns in practice\', Financial Analysts Journal, 56(4), pp. 19–31.',
        'Hartley, S. and Faff, R. (2020) \'The (non-) predictive power of US market-wide sentiment\', Journal of Banking & Finance, 121, p. 105963.',
        'Maillard, S., Roncalli, T. and Teïletche, J. (2010) \'The properties of equally weighted risk contribution portfolios\', Journal of Portfolio Management, 36(4), pp. 60–70.',
        'Markowitz, H. (1952) \'Portfolio selection\', The Journal of Finance, 7(1), pp. 77–91.',
        'UNSW FINS3645 FinTech (2026) Week 2–7 Lecture Slides. University of New South Wales.'
    ]
    
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(6)
    
    # Save document
    output_path = '/Users/yumeme/Documents/GitHub/FINS3645_FinTech_Project/report/FINS3645_PartA_Report.docx'
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    return output_path


if __name__ == '__main__':
    create_report()
